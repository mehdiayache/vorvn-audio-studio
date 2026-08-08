#!/usr/bin/env python3
"""
Turn a dropped document into plain text, plus a sensible split into blocks.

Handles .txt/.md, .docx and .pdf. Everything else is rejected with a message
naming what is supported, rather than producing mojibake.
"""

import io
import re

SUPPORTED = (".txt", ".md", ".markdown", ".text", ".docx", ".pdf")

# One block per paragraph is too fine for prose and too coarse for dialogue, so
# paragraphs are the unit and long ones get split at sentence ends.
BLOCK_TARGET = 480


def extract(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith((".txt", ".md", ".markdown", ".text")):
        return _from_text(data)
    raise ValueError(
        f"Can't read '{filename}'. Supported: {', '.join(SUPPORTED)}."
    )


def _from_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _from_docx(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Tables hold real script content often enough to be worth pulling in.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" ".join(cells))
    return "\n\n".join(p for p in parts if p.strip())


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        try:
            reader.decrypt("")   # many PDFs are "encrypted" with an empty password
        except Exception:
            raise ValueError("That PDF is password protected.")
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(p for p in pages if p)
    if not text.strip():
        raise ValueError(
            "No text in that PDF — it's probably scanned images, which need OCR."
        )
    return text


def clean(text: str) -> str:
    """Undo the things that make extracted text read badly aloud."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"-\n(?=[a-z])", "", text)       # de-hyphenate line-wrapped words
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)          # collapse big gaps
    text = re.sub(r"(?<![\n.!?:;])\n(?![\n])", " ", text)  # unwrap soft line breaks
    return text.strip()


def to_blocks(text: str, target: int = BLOCK_TARGET) -> list:
    """Split cleaned text into block-sized pieces.

    Paragraph breaks are the natural seam, but honouring every one of them turns
    dialogue into hundreds of one-line blocks. So short paragraphs are packed
    together up to the target and over-long ones are split at sentence ends.
    """
    pieces = []
    for paragraph in (p.strip() for p in re.split(r"\n+", clean(text))):
        if not paragraph:
            continue
        if len(paragraph) <= target:
            pieces.append(paragraph)
            continue
        current = ""
        for sentence in re.split(r"(?<=[.!?])\s+", paragraph):
            if current and len(current) + len(sentence) + 1 > target:
                pieces.append(current.strip())
                current = sentence
            else:
                current = f"{current} {sentence}".strip()
        if current:
            pieces.append(current.strip())

    blocks = []
    for piece in pieces:
        if blocks and len(blocks[-1]) + len(piece) + 1 <= target:
            blocks[-1] = f"{blocks[-1]}\n{piece}"
        else:
            blocks.append(piece)
    return blocks
